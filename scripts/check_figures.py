#!/usr/bin/env python3
"""
check_figures.py - Arithmetic and consistency checks for documents containing figures.

Two modes:

  check    Verify tables inside one or more documents. Confirms that total rows equal
           the sum of their columns and that derived "change" columns equal the
           difference they claim. Also lists, with line numbers, any figure of 1,000 or
           more that appears in prose but matches no value in any table, so those can
           be reconciled by hand.

  compare  Diff figures between two versions of a document, keyed on table row labels.
           Surfaces values that changed between drafts without anyone announcing it.

Usage:
    python3 check_figures.py check draft.md
    python3 check_figures.py check v1.md v2.md
    python3 check_figures.py compare v1.md v2.md

Reads markdown and plain text with no dependencies. Reads .docx and .pdf as well when
python-docx or pdfplumber happen to be installed, and says so if they are not.

Exit code is 1 when any check fails, so this can gate a workflow.

What this cannot do: reconcile a prose claim against a subset of table rows. "$230K in
other tooling" versus a table whose relevant rows sum to $260K requires knowing which
rows "other tooling" refers to. That judgment stays with the reader. See
references/numeric-integrity.md for the checks that need a human or an agent.
"""

import re
import sys
from pathlib import Path

MONEY = re.compile(r'\(?\$\s?-?[\d,]+(?:\.\d+)?\s?[KMB]?\)?|\(?-?[\d,]+(?:\.\d+)?\s?%\)?')
TOTAL_LABEL = re.compile(r'\b(?:total|subtotal|net\s+total|combined\s+total|overall\s+total)\b', re.I)
DERIVED_LABEL = re.compile(r'\b(change|delta|variance|diff|increase|decrease|impact)\b', re.I)
SUFFIX = {'K': 1_000, 'M': 1_000_000, 'B': 1_000_000_000}


def parse_number(cell):
    """Return a float for a numeric cell, or None. Parentheses mean negative."""
    if cell is None:
        return None
    s = str(cell).strip().replace('*', '').replace('`', '')
    if not s or s in {'-', '--', 'n/a', 'N/A', 'TBD'}:
        return None
    negative = s.startswith('(') and s.endswith(')')
    s = s.strip('()').strip()
    s = s.replace('$', '').replace(',', '').replace('+', '').strip()
    is_pct = s.endswith('%')
    s = s.rstrip('%').strip()
    mult = 1
    if s and s[-1].upper() in SUFFIX:
        mult = SUFFIX[s[-1].upper()]
        s = s[:-1].strip()
    try:
        value = float(s) * mult
    except ValueError:
        return None
    if negative:
        value = -value
    return ('pct', value) if is_pct else ('num', value)


def numeric(cell):
    parsed = parse_number(cell)
    return parsed[1] if parsed else None


def read_text(path):
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == '.docx':
        try:
            import docx
        except ImportError:
            sys.exit(f"{path}: reading .docx needs python-docx. Convert to markdown first.")
        doc = docx.Document(str(p))
        lines = [par.text for par in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                lines.append('| ' + ' | '.join(c.text.strip() for c in row.cells) + ' |')
        return '\n'.join(lines)
    if suffix == '.pdf':
        try:
            import pdfplumber
        except ImportError:
            sys.exit(f"{path}: reading .pdf needs pdfplumber. Convert to markdown first.")
        out = []
        with pdfplumber.open(str(p)) as pdf:
            for page in pdf.pages:
                out.append(page.extract_text() or '')
                for table in page.extract_tables():
                    for row in table:
                        cells = [(c or '').replace('\n', ' ').strip() for c in row]
                        if any(cells):
                            out.append('| ' + ' | '.join(cells) + ' |')
        return '\n'.join(out)
    return p.read_text(encoding='utf-8', errors='replace')


def extract_tables(text):
    """Return [{'header': [...], 'rows': [(lineno, [cells])], 'start': lineno}]."""
    tables, current = [], None
    for i, line in enumerate(text.splitlines(), 1):
        stripped = line.strip()
        is_row = stripped.startswith('|') and stripped.count('|') >= 2
        if is_row:
            cells = [c.strip().replace('*', '') for c in stripped.strip('|').split('|')]
            cells = [c for c in cells]
            if set(''.join(cells).replace(' ', '')) <= {'-', ':'} and cells:
                continue  # separator row
            if current is None:
                current = {'header': cells, 'rows': [], 'start': i}
            else:
                current['rows'].append((i, cells))
        else:
            if current and current['rows']:
                tables.append(current)
            current = None
    if current and current['rows']:
        tables.append(current)
    return tables


def check_table(table, findings, path):
    header = table['header']
    rows = table['rows']
    width = max([len(header)] + [len(r[1]) for r in rows])

    data_rows = [(ln, c) for ln, c in rows if not TOTAL_LABEL.search(c[0] if c else '')]
    total_rows = [(ln, c) for ln, c in rows if TOTAL_LABEL.search(c[0] if c else '')]

    # 1. Total rows must equal the sum of the data rows above them, column by column.
    for lineno, total_cells in total_rows:
        for col in range(1, width):
            stated = numeric(total_cells[col]) if col < len(total_cells) else None
            if stated is None:
                continue
            parts = [numeric(c[col]) for ln, c in data_rows if ln < lineno and col < len(c)]
            parts = [p for p in parts if p is not None]
            if len(parts) < 2:
                continue
            actual = sum(parts)
            if abs(actual - stated) > 0.51:
                col_name = header[col] if col < len(header) else f'column {col + 1}'
                findings.append(
                    f"{path}:{lineno}  TOTAL MISMATCH in '{col_name}': "
                    f"stated {fmt(stated)}, rows sum to {fmt(actual)} "
                    f"(off by {fmt(actual - stated)})"
                )

    # 2. A column named change/delta/variance should equal the column before it minus
    #    the first numeric column. This is the shape almost every spend table uses.
    first_num_col = None
    for col in range(1, width):
        if any(numeric(c[col]) is not None for _, c in rows if col < len(c)):
            first_num_col = col
            break
    if first_num_col is None:
        return
    for col in range(1, width):
        name = header[col] if col < len(header) else ''
        if not DERIVED_LABEL.search(name) or col - 1 <= first_num_col - 1:
            continue
        for lineno, cells in rows:
            if col >= len(cells):
                continue
            stated = numeric(cells[col])
            current = numeric(cells[col - 1]) if col - 1 < len(cells) else None
            base = numeric(cells[first_num_col]) if first_num_col < len(cells) else None
            if None in (stated, current, base):
                continue
            expected = current - base
            if abs(expected - stated) > 0.51:
                label = cells[0] or f'row at line {lineno}'
                findings.append(
                    f"{path}:{lineno}  DERIVED COLUMN MISMATCH in '{name}' for '{label}': "
                    f"stated {fmt(stated)}, {fmt(current)} minus {fmt(base)} is {fmt(expected)}"
                )


def fmt(v):
    if v is None:
        return 'n/a'
    if abs(v - round(v)) < 0.005:
        return f"{'-' if v < 0 else ''}{abs(int(round(v))):,}"
    return f"{v:,.2f}"


def row_index(text):
    """Map normalized table row labels to the values on that row."""
    index = {}
    for table in extract_tables(text):
        for lineno, cells in table['rows']:
            if not cells or not cells[0]:
                continue
            key = re.sub(r'[^a-z0-9]+', ' ', cells[0].lower()).strip()
            if not key:
                continue
            values = [numeric(c) for c in cells[1:]]
            values = [v for v in values if v is not None]
            if values:
                index.setdefault(key, []).append((lineno, values))
    return index


def inventory(text):
    """Every currency and percentage figure in the document, by line."""
    found = []
    for i, line in enumerate(text.splitlines(), 1):
        for match in MONEY.finditer(line):
            parsed = parse_number(match.group())
            if parsed:
                found.append((i, match.group().strip(), parsed[0], parsed[1]))
    return found


def cmd_check(paths):
    findings, review = [], []
    for path in paths:
        text = read_text(path)
        tables = extract_tables(text)
        print(f"\n{path}: {len(tables)} table(s), {len(inventory(text))} figure(s)")
        for table in tables:
            check_table(table, findings, path)

        # Figures that appear in prose but in no table. Not errors. These are the
        # ones a reader has to reconcile by hand, because deciding whether "$230K in
        # other tooling" matches a table is a question about which rows it refers to.
        figs = inventory(text)
        in_tables = set()
        for t in tables:
            in_tables.add(t['start'])
            for ln, _ in t['rows']:
                in_tables.add(ln)
        table_vals = {v for ln, _, kind, v in figs if ln in in_tables and kind == 'num'}
        for ln, raw, kind, v in figs:
            if ln in in_tables or kind != 'num' or abs(v) < 1000:
                continue
            if v not in table_vals:
                review.append(f"{path}:{ln}  {raw}")

    print()
    status = 0
    if findings:
        print(f"ISSUES ({len(findings)}) - arithmetic that does not hold:\n")
        for f in findings:
            print(f"  {f}")
        print("\n  Re-add every column after changing any row. A total carried over")
        print("  from a previous draft is the most common error in these documents.")
        status = 1
    else:
        print("ISSUES: none. Every total and derived column checks out.")

    if review:
        print(f"\nREVIEW ({len(review)}) - figures in prose that appear in no table.")
        print("Not errors. Reconcile each one against the table by hand:\n")
        for r in review:
            print(f"  {r}")

    print("\nChecks needing judgment rather than arithmetic are listed in")
    print("references/numeric-integrity.md. Run those too.")
    return status


def leading_tokens(key):
    return key.split()


def best_match(key, candidates):
    """Pair row labels across drafts by shared leading words.

    Labels get shortened between versions: 'Splunk Cloud spend reduced under this
    renewal' becomes 'Splunk Cloud'. Exact matching would report that as one row
    deleted and another added, hiding the fact that its value moved.
    """
    if key in candidates:
        return key
    target = leading_tokens(key)
    scored = []
    for cand in candidates:
        toks = leading_tokens(cand)
        n = 0
        for a, b in zip(target, toks):
            if a != b:
                break
            n += 1
        if n:
            scored.append((n, cand))
    if not scored:
        return None
    scored.sort(reverse=True)
    if len(scored) > 1 and scored[0][0] == scored[1][0]:
        return None  # ambiguous, better to say nothing
    return scored[0][1]


def cmd_compare(old_path, new_path):
    old, new = row_index(read_text(old_path)), row_index(read_text(new_path))
    changed, added, removed = [], [], []
    matched_old = set()
    for key, entries in new.items():
        match = best_match(key, old)
        if match is None:
            added.append(key)
            continue
        matched_old.add(match)
        old_vals = old[match][0][1]
        new_vals = entries[0][1]
        # Compare only the overlapping leading values. Extra columns in the newer
        # draft (a Year 2 and Year 3 added, say) are a structural change, not a
        # changed figure, and flagging them would bury the values that really moved.
        span = min(len(old_vals), len(new_vals))
        if old_vals[:span] != new_vals[:span]:
            label = key if match == key else f"{match} -> {key}"
            changed.append((label, old_vals[:span], new_vals[:span], entries[0][0]))
    for key in old:
        if key not in matched_old:
            removed.append(key)

    print(f"\nComparing {old_path} to {new_path}\n")
    if changed:
        print("VALUES CHANGED between versions:")
        for key, o, n, ln in changed:
            print(f"  line {ln}  '{key}'")
            print(f"      was: {', '.join(fmt(v) for v in o)}")
            print(f"      now: {', '.join(fmt(v) for v in n)}")
        print()
    if removed:
        print("ROWS REMOVED: " + ', '.join(f"'{k}'" for k in removed))
        print("  Check that any total which included these rows was recalculated.\n")
    if added:
        print("ROWS ADDED: " + ', '.join(f"'{k}'" for k in added) + "\n")
    if not (changed or removed or added):
        print("No table rows changed.\n")
        return 0
    print("Every changed figure needs a reason. A value that moved between drafts")
    print("without anyone noticing is the one that gets challenged in the meeting.")
    return 1


def main():
    args = sys.argv[1:]
    if not args or args[0] in {'-h', '--help'}:
        print(__doc__)
        return 0
    mode, rest = args[0], args[1:]
    if mode == 'check' and rest:
        return cmd_check(rest)
    if mode == 'compare' and len(rest) == 2:
        return cmd_compare(rest[0], rest[1])
    print(__doc__)
    return 2


if __name__ == '__main__':
    sys.exit(main())
